# import pyaudio
import wave
import os
from termcolor import colored
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
from docx import Document
from backend.key_points_service import key_points_extraction
from backend.action_items_service import action_item_extraction

def transcribe_chunk(model, audio):
    print('trascribing...')
    segments, _ = model.transcribe(audio, beam_size=5)
    final = ""
    for segment in segments:
        final += segment.text
    # print('trascribed...')
    return final


def record_chunk(audio_buffer, chunk_file):
    # return frames    
    wf = wave.open(chunk_file, 'wb' )
    wf.setnchannels(1)
    wf.setsampwidth(4)
    wf.setframerate(16000)
    wf.writeframes(b''+audio_buffer)
    wf.close()
    # with open(chunk_file, 'wb') as f:
    #     f.write(audio_buffer)
    # print('file written...')
    
def get_transcipts(model, audio_buffer):
    chunk_file = "output.wav"
    record_chunk(audio_buffer, chunk_file)
    transcription = transcribe_chunk(model, chunk_file)
    #  (transcription)
    print(colored(transcription, color="green"))
    # os.remove(chunk_file)
    # accumulated_transcription += transcription + " "
    return transcription

def abstract_summary_extraction(text, max_output_length=500):
    from transformers import AutoModelForSeq2SeqLM, AutoTokenizer
    model_path= "models/transcript-summerizer/"
    device = "cuda" if torch.cuda.is_available() else "cpu"
    
    tokenizer = AutoTokenizer.from_pretrained(model_path)

    model = AutoModelForSeq2SeqLM.from_pretrained(model_path).to(device)
    # Preprocess the text
    input_encodings = tokenizer(text, return_tensors="pt", truncation=True, padding="max_length")

    # Generate the summary
    summary_ids = model.generate(input_encodings['input_ids'].to(device),
                                 attention_mask=input_encodings['attention_mask'].to(device),
                                 max_length=max_output_length,
                                 num_beams=6,  # Increased beam width
                                 temperature=2,  # Adjusted temperature
                                 repetition_penalty=1,  # Increased repetition penalty
                                 length_penalty=10,  # Adjust length penalty
                                 early_stopping=True)

    # Decode the summary
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

    return summary


def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)


def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items_summary': action_items
    }

def generate_minutes(file_name):
    
    with open(file_name, "r") as file:
        transcription = file.read()
    
    minutes = meeting_minutes(transcription)
    print(minutes)
    save_as_docx(minutes, 'meeting_minutes.docx')
    return minutes